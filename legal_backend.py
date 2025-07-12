import os
import pytesseract
from pdf2image import convert_from_path
from docx import Document as DocxDocument
from datetime import datetime
import csv

from llama_index.core import Document as LlamaDocument
from llama_index.core import VectorStoreIndex, Settings
from llama_index.llms.ollama import Ollama
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# ✅ Setup: LLM + Embedding
llm = Ollama(model="llama3", request_timeout=120)
embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")
Settings.llm = llm
Settings.embed_model = embed_model

chat_history = []


def load_document(file_path):
    text = ""
    try:
        if file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        elif file_path.endswith(".docx"):
            doc = DocxDocument(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])

        elif file_path.endswith(".pdf"):
            try:
                from llama_index.readers.file import PDFReader
                reader = PDFReader()
                pdf_docs = reader.load_data(file_path)
                text = "\n".join([doc.text for doc in pdf_docs])
            except Exception:
                images = convert_from_path(file_path)
                for img in images:
                    text += pytesseract.image_to_string(img)

        if not text.strip():
            raise ValueError(
                "❌ No extractable text found in the uploaded document.")

        return [LlamaDocument(text=text)]

    except Exception as e:
        raise RuntimeError(f"Failed to load document: {e}")


def build_index(documents):
    return VectorStoreIndex.from_documents(documents)


def save_to_log(filename, category, content):
    os.makedirs("logs", exist_ok=True)
    log_path = os.path.join("logs", "session_log.csv")
    header = ["timestamp", "filename", "category", "content"]

    if not os.path.exists(log_path):
        with open(log_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(header)

    with open(log_path, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([datetime.now().isoformat(),
                        filename, category, content])


def summarize_document(documents):
    if not documents:
        return "⚠️ No document to summarize. Please upload a valid file."

    text = documents[0].text.strip()
    if not text:
        return "⚠️ Document is empty."

    short_text = text[:12000]  # approx first 3000 tokens

    prompt = f"""
You are a legal assistant. Summarize the following legal document using markdown headings and bullet points.
Be clear, concise, and highlight key clauses, parties involved, and obligations.

Document:
{short_text}

Summary:
"""
    try:
        print("🔍 Sending prompt to LLM...")
        summary = llm.complete(prompt).text.strip()
        if not summary:
            print("⚠️ Empty summary returned.")
            return "⚠️ The AI returned an empty summary. Try again or check the document content."
        print("✅ Summary received.")
        save_to_log("uploaded", "summary", summary)
        return summary
    except Exception as e:
        print("❌ Summarization Error:", e)
        return f"❌ Summarization failed: {e}"


def highlight_clauses(documents):
    if not documents:
        return "⚠️ No document to analyze. Please upload a valid file."
    index = build_index(documents)
    response = index.as_query_engine().query(
        "List the important legal clauses and tag them under categories like Confidentiality, Termination, Liability, Dispute Resolution, etc. Explain each in plain language."
    )
    clauses = str(response)
    save_to_log("uploaded", "highlighted_clauses", clauses)
    return clauses


def clause_breakdown(documents):
    if not documents:
        return "⚠️ No document available for clause breakdown."
    index = build_index(documents)
    response = index.as_query_engine().query(
        "Break this legal document into individual clauses and explain each one clearly."
    )
    breakdown = str(response)
    save_to_log("uploaded", "clause_breakdown", breakdown)
    return breakdown


def simplify_legal_jargon(documents):
    if not documents:
        return "⚠️ No document to simplify."
    index = build_index(documents)
    response = index.as_query_engine().query(
        "Rewrite this legal document in extremely simple, everyday language that anyone can understand."
    )
    simplified = str(response)
    save_to_log("uploaded", "simplified", simplified)
    return simplified


def answer_query(documents, query):
    global chat_history
    combined_text = "\n".join(
        [doc.text for doc in documents if doc.text.strip()])
    if not combined_text.strip():
        combined_text = "No usable text found in the uploaded document."

    history = "\n".join([f"User: {q}\nAI: {a}" for q, a in chat_history])

    prompt = f"""
You are a helpful AI with expertise in legal and general questions.
Always answer clearly, even if the question is not related to any document.

Uploaded Document:
{combined_text[:2000]}

Chat History:
{history}

User: {query}
AI:"""

    response = llm.complete(prompt).text.strip()
    chat_history.append((query, response))
    save_to_log("uploaded", "qa", f"Q: {query}\nA: {response}")
    return response


def extract_entities(documents):
    if not documents:
        return "⚠️ No document to extract entities from."
    index = build_index(documents)
    response = index.as_query_engine().query(
        "Extract all named entities from this legal document. Categorize them into: People, Organizations, Dates, Locations, Legal Terms."
    )
    entities = str(response)
    save_to_log("uploaded", "entities", entities)
    return entities


def export_highlighted_pdf(documents, original_file_path):
    from fpdf import FPDF
    if not documents:
        return None
    text = documents[0].text
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)
    os.makedirs("outputs", exist_ok=True)
    output_path = os.path.join("outputs", "highlighted_output.pdf")
    pdf.output(output_path)
    save_to_log("uploaded", "exported_pdf", "Generated highlighted PDF.")
    return output_path


def compare_documents(doc1, doc2):
    if not doc1 or not doc2:
        return "⚠️ Both documents must be uploaded for comparison."

    index = build_index(doc1 + doc2)
    prompt = """
Compare the two legal documents provided. Highlight:
- Key similarities and differences in clauses
- Any mismatched obligations or terms
- Differences in parties, durations, dispute resolution, liabilities, etc.
Use clear headings and bullet points.
"""
    response = index.as_query_engine().query(prompt)
    save_to_log("uploaded", "comparison", str(response))
    return str(response)
